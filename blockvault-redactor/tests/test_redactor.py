"""Tests for BlockVault Redactor."""
from __future__ import annotations

import io
import json

import pytest
from fastapi.testclient import TestClient
import fitz  # PyMuPDF
from docx import Document

from app.main import app
from app.detector import detect_entities


client = TestClient(app)

# ---------------------------------------------------------------------------
# Test entity detection
# ---------------------------------------------------------------------------

def test_regex_detection():
    """Test standard regex patterns (SSN, Email, Phone, Bank)."""
    text = "Contact john.doe@example.com or 555-123-4567. SSN is 838-23-1111. Account: GB82WEST12345698765432"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "EMAIL" in types
    assert "PHONE" in types
    assert "SSN" in types
    assert "BANK_ACCOUNT" in types


def test_spacy_detection():
    """Test spaCy NER (PERSON, ORG)."""
    text = "Elon Musk recently bought Twitter which is based in San Francisco."
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "PERSON" in types
    assert "ORG" in types


def test_deduplication():
    """Test overlapping entity resolution."""
    # A long alphanumeric string might match both ID_NUMBER and something else
    text = "My ID number is AB12345678X"
    entities = detect_entities(text)
    # Should resolve to a single entity for that span
    spans = [(e.start, e.end) for e in entities]
    assert len(spans) == len(set(spans))


# ---------------------------------------------------------------------------
# Test PDF endpoints
# ---------------------------------------------------------------------------

def _create_test_pdf() -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Hello John Smith. My number is 555-999-0000.")
    out = io.BytesIO()
    doc.save(out)
    doc.close()
    return out.getvalue()


def test_analyze_pdf():
    pdf_bytes = _create_test_pdf()
    response = client.post(
        "/analyze",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")}
    )
    assert response.status_code == 200
    data = response.json()
    assert "entities" in data
    
    types = {e["entity_type"] for e in data["entities"]}
    assert "PERSON" in types
    assert "PHONE" in types

    # Check bounding box
    bbox_ent = next(e for e in data["entities"] if e["entity_type"] == "PERSON")
    assert "bbox" in bbox_ent
    assert "page" in bbox_ent
    assert bbox_ent["page"] == 1


def test_redact_pdf():
    pdf_bytes = _create_test_pdf()
    
    # Analyze first
    analyze_resp = client.post(
        "/analyze",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")}
    )
    entities = analyze_resp.json()["entities"]

    # Then redact
    redact_resp = client.post(
        "/redact",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
        data={"entities": json.dumps(entities)}
    )
    assert redact_resp.status_code == 200
    assert redact_resp.headers["Content-Type"] == "application/pdf"
    
    # Verify redaction: text should be gone
    doc = fitz.open(stream=redact_resp.content, filetype="pdf")
    text = doc[0].get_text()
    assert "John Smith" not in text
    assert "555-999-0000" not in text
    doc.close()


# ---------------------------------------------------------------------------
# Test DOCX endpoints
# ---------------------------------------------------------------------------

def _create_test_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Alice lives in New York.")
    doc.add_paragraph("Her email is alice@company.com")
    # Add a table to test table traversal
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Secret project details"
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def test_analyze_docx():
    docx_bytes = _create_test_docx()
    response = client.post(
        "/analyze",
        files={"file": ("test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response.status_code == 200
    data = response.json()
    
    types = {e["entity_type"] for e in data["entities"]}
    assert "PERSON" in types  # Alice
    assert "ADDRESS" in types  # New York
    assert "EMAIL" in types  # alice@company.com


def test_redact_docx():
    docx_bytes = _create_test_docx()
    
    analyze_resp = client.post(
        "/analyze",
        files={"file": ("test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    entities = analyze_resp.json()["entities"]

    redact_resp = client.post(
        "/redact",
        files={"file": ("test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"entities": json.dumps(entities)}
    )
    assert redact_resp.status_code == 200
    
    # Reload redacted document
    doc = Document(io.BytesIO(redact_resp.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    
    assert "Alice" not in text
    assert "New York" not in text
    assert "alice@company.com" not in text
    assert "[REDACTED]" in text


# ---------------------------------------------------------------------------
# Security & Config Tests
# ---------------------------------------------------------------------------

def test_unsupported_file():
    txt_bytes = b"Hello world"
    response = client.post(
        "/analyze",
        files={"file": ("test.txt", txt_bytes, "text/plain")}
    )
    assert response.status_code == 400
    assert "Unsupported file type" in response.json()["detail"]


def test_metadata_cleaning(monkeypatch):
    """Ensure DOCX metadata is stripped during redaction."""
    docx_bytes = _create_test_docx()
    # Add metadata to original
    doc = Document(io.BytesIO(docx_bytes))
    doc.core_properties.author = "Sneaky Attacker"
    doc.core_properties.comments = "Hidden tracking code"
    
    out = io.BytesIO()
    doc.save(out)
    dirty_bytes = out.getvalue()
    
    redact_resp = client.post(
        "/redact",
        files={"file": ("test.docx", dirty_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"entities": json.dumps([])}  # Pass empty entities to just trigger cleaning
    )
    
    clean_doc = Document(io.BytesIO(redact_resp.content))
    assert clean_doc.core_properties.author == ""
    assert clean_doc.core_properties.comments == ""
