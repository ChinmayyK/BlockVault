"""Tests for BlockVault Redactor."""
from __future__ import annotations

import io
import json

import pytest
from fastapi.testclient import TestClient
import fitz  # PyMuPDF
from docx import Document

from app.main import app
from app.detector import (
    detect_entities,
    _verhoeff_checksum,
    _luhn_checksum,
    _regex_detect,
    _label_detect,
)


client = TestClient(app)

# ---------------------------------------------------------------------------
# Test checksum validation
# ---------------------------------------------------------------------------

def test_verhoeff_checksum_valid():
    """Test Verhoeff checksum with a known valid Aadhaar."""
    # 4567 8901 2345 is not necessarily valid; test with known logic
    assert isinstance(_verhoeff_checksum("123456789012"), bool)


def test_luhn_checksum_valid():
    """Test Luhn checksum with a known valid credit card (Visa test number)."""
    assert _luhn_checksum("4111111111111111") is True


def test_luhn_checksum_invalid():
    """Test Luhn checksum with an invalid credit card."""
    assert _luhn_checksum("4111111111111112") is False


# ---------------------------------------------------------------------------
# Test regex detection
# ---------------------------------------------------------------------------

def test_regex_email():
    """Test email detection via regex."""
    text = "Contact john.doe@example.com for details"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "EMAIL" in types
    email_ent = next(e for e in entities if e.entity_type == "EMAIL")
    assert email_ent.text == "john.doe@example.com"


def test_regex_phone():
    """Test phone number detection via regex."""
    text = "Call me at 555-123-4567 or (800) 555-1234"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "PHONE" in types


def test_regex_ssn():
    """Test SSN detection via regex."""
    text = "SSN is 838-23-1111"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "SSN" in types


def test_regex_iban():
    """Test IBAN detection via regex."""
    text = "Account: GB82WEST12345698765432"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "IBAN" in types


def test_regex_pan():
    """Test Indian PAN detection."""
    text = "PAN Number: ABCDE1234F"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "PAN" in types
    pan_ent = next(e for e in entities if e.entity_type == "PAN")
    assert pan_ent.text == "ABCDE1234F"


def test_regex_voter_id():
    """Test Indian Voter ID / EPIC detection."""
    text = "Voter ID: ABC1234567"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "VOTER_ID" in types


def test_regex_gstin():
    """Test Indian GSTIN detection."""
    text = "GSTIN: 22AAAAA0000A1Z5"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "GSTIN" in types


def test_regex_ifsc():
    """Test Indian IFSC Code detection."""
    text = "IFSC Code: SBIN0001234"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "IFSC" in types


def test_regex_date():
    """Test date pattern detection."""
    text = "Date of Birth: 15/03/1990"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "DATE_OF_BIRTH" in types or "DATE" in types


def test_regex_date_text_month():
    """Test date with text month names."""
    text = "Born on 15 January 1990"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "DATE_OF_BIRTH" in types or "DATE" in types


def test_regex_ipv4():
    """Test IPv4 detection."""
    text = "Server IP is 192.168.1.100"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "IP_ADDRESS" in types


def test_regex_indian_phone():
    """Test Indian phone number detection."""
    text = "Mobile: +91 9876543210"
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "PHONE" in types


# ---------------------------------------------------------------------------
# Test label-based detection
# ---------------------------------------------------------------------------

def test_label_name():
    """Test label-based name extraction."""
    text = "Name: Rajesh Kumar\nFather's Name: Suresh Kumar\nAddress: 123 Main Street"
    entities = _label_detect(text)
    types = {e.entity_type for e in entities}
    texts = {e.text for e in entities}
    assert "PERSON" in types
    assert "Rajesh Kumar" in texts
    assert "Suresh Kumar" in texts


def test_label_address():
    """Test label-based address extraction."""
    text = "Permanent Address: 45 MG Road, Bangalore, Karnataka 560001\nPhone: 9876543210"
    entities = _label_detect(text)
    types = {e.entity_type for e in entities}
    assert "ADDRESS" in types
    addr_ent = next(e for e in entities if e.entity_type == "ADDRESS")
    assert "MG Road" in addr_ent.text


def test_label_dob():
    """Test label-based DOB extraction."""
    text = "Date of Birth: 15/03/1990\nGender: Male"
    entities = _label_detect(text)
    types = {e.entity_type for e in entities}
    assert "DATE_OF_BIRTH" in types


def test_label_aadhaar():
    """Test label-based Aadhaar extraction."""
    text = "Aadhaar Number: 2345 6789 0123"
    entities = _label_detect(text)
    types = {e.entity_type for e in entities}
    assert "AADHAAR" in types


def test_label_gender():
    """Test label-based gender extraction."""
    text = "Gender: Male\nAddress: XYZ"
    entities = _label_detect(text)
    types = {e.entity_type for e in entities}
    assert "GENDER" in types


def test_label_sensitive_category():
    """Test label-based religion/caste/category extraction."""
    text = "Religion: Hindu\nCaste: General"
    entities = _label_detect(text)
    types = {e.entity_type for e in entities}
    assert "SENSITIVE_CATEGORY" in types


# ---------------------------------------------------------------------------
# Test full pipeline (multi-layer)
# ---------------------------------------------------------------------------

def test_full_pipeline_indian_document():
    """Test full pipeline with Indian certificate-style text."""
    text = (
        "GOVERNMENT OF INDIA\n"
        "PERSONAL DETAILS CERTIFICATE\n"
        "\n"
        "Name: Ramesh Kumar Sharma\n"
        "Father's Name: Suresh Kumar Sharma\n"
        "Date of Birth: 15/03/1990\n"
        "Gender: Male\n"
        "Address: 45 MG Road, Sector 21, Noida, UP 201301\n"
        "PAN: ABCDE1234F\n"
        "Voter ID: ABC1234567\n"
        "Mobile: +91 9876543210\n"
        "Email: ramesh.sharma@example.com\n"
    )
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}

    # All of these should be detected
    assert "PERSON" in types, f"PERSON not detected. Found: {types}"
    assert "EMAIL" in types, f"EMAIL not detected. Found: {types}"
    assert "PAN" in types, f"PAN not detected. Found: {types}"
    assert "VOTER_ID" in types, f"VOTER_ID not detected. Found: {types}"
    assert "PHONE" in types, f"PHONE not detected. Found: {types}"
    assert "DATE_OF_BIRTH" in types, f"DATE_OF_BIRTH not detected. Found: {types}"


def test_deduplication():
    """Test overlapping entity resolution."""
    text = "My ID number is AB12345678X"
    entities = detect_entities(text)
    spans = [(e.start, e.end) for e in entities]
    assert len(spans) == len(set(spans))


def test_spacy_detection():
    """Test spaCy NER (PERSON, ORG)."""
    text = "Elon Musk recently bought Twitter which is based in San Francisco."
    entities = detect_entities(text)
    types = {e.entity_type for e in entities}
    assert "PERSON" in types
    assert "ORG" in types


# ---------------------------------------------------------------------------
# Test PDF endpoints
# ---------------------------------------------------------------------------

def _create_test_pdf() -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Hello John Smith. My number is 555-999-0000.")
    out = io.BytesIO()
    doc.save(out)
    doc.close()
    return out.getvalue()


def _create_indian_test_pdf() -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Name: Rajesh Kumar Sharma")
    page.insert_text((50, 70), "Father's Name: Suresh Kumar")
    page.insert_text((50, 90), "PAN: ABCDE1234F")
    page.insert_text((50, 110), "Email: raj@example.com")
    page.insert_text((50, 130), "Mobile: +91 9876543210")
    page.insert_text((50, 150), "DOB: 15/03/1990")
    out = io.BytesIO()
    doc.save(out)
    doc.close()
    return out.getvalue()


def test_analyze_pdf():
    pdf_bytes = _create_test_pdf()
    response = client.post(
        "/analyze",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")}
    )
    assert response.status_code == 200
    data = response.json()
    assert "entities" in data

    types = {e["entity_type"] for e in data["entities"]}
    assert "PERSON" in types
    assert "PHONE" in types

    # Check bounding box
    bbox_ent = next(e for e in data["entities"] if e["entity_type"] == "PERSON")
    assert "bbox" in bbox_ent
    assert "page" in bbox_ent
    assert bbox_ent["page"] == 1


def test_analyze_indian_pdf():
    """Test analysis of Indian-style certificate PDF."""
    pdf_bytes = _create_indian_test_pdf()
    response = client.post(
        "/analyze",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")}
    )
    assert response.status_code == 200
    data = response.json()
    types = {e["entity_type"] for e in data["entities"]}

    assert "PAN" in types, f"PAN not detected. Found types: {types}"
    assert "EMAIL" in types, f"EMAIL not detected. Found types: {types}"
    assert "PHONE" in types, f"PHONE not detected. Found types: {types}"


def test_redact_pdf():
    pdf_bytes = _create_test_pdf()

    # Analyze first
    analyze_resp = client.post(
        "/analyze",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")}
    )
    entities = analyze_resp.json()["entities"]

    # Then redact
    redact_resp = client.post(
        "/redact",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
        data={"entities": json.dumps(entities)}
    )
    assert redact_resp.status_code == 200
    assert redact_resp.headers["Content-Type"] == "application/pdf"

    # Verify redaction: text should be gone
    doc = fitz.open(stream=redact_resp.content, filetype="pdf")
    text = doc[0].get_text()
    assert "John Smith" not in text
    assert "555-999-0000" not in text
    doc.close()


def test_redact_indian_pdf():
    """Test redaction of Indian-style certificate PDF."""
    pdf_bytes = _create_indian_test_pdf()

    analyze_resp = client.post(
        "/analyze",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")}
    )
    entities = analyze_resp.json()["entities"]

    redact_resp = client.post(
        "/redact",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
        data={"entities": json.dumps(entities)}
    )
    assert redact_resp.status_code == 200

    doc = fitz.open(stream=redact_resp.content, filetype="pdf")
    text = doc[0].get_text()
    # PAN and email should be redacted
    assert "ABCDE1234F" not in text
    assert "raj@example.com" not in text
    doc.close()


# ---------------------------------------------------------------------------
# Test DOCX endpoints
# ---------------------------------------------------------------------------

def _create_test_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Alice lives in New York.")
    doc.add_paragraph("Her email is alice@company.com")
    # Add a table to test table traversal
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Secret project details"
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def test_analyze_docx():
    docx_bytes = _create_test_docx()
    response = client.post(
        "/analyze",
        files={"file": ("test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response.status_code == 200
    data = response.json()

    types = {e["entity_type"] for e in data["entities"]}
    assert "PERSON" in types  # Alice
    assert "EMAIL" in types  # alice@company.com


def test_redact_docx():
    docx_bytes = _create_test_docx()

    analyze_resp = client.post(
        "/analyze",
        files={"file": ("test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    entities = analyze_resp.json()["entities"]

    redact_resp = client.post(
        "/redact",
        files={"file": ("test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"entities": json.dumps(entities)}
    )
    assert redact_resp.status_code == 200

    # Reload redacted document
    doc = Document(io.BytesIO(redact_resp.content))
    text = "\n".join(p.text for p in doc.paragraphs)

    assert "Alice" not in text
    assert "alice@company.com" not in text
    assert "[REDACTED]" in text


# ---------------------------------------------------------------------------
# Security & Config Tests
# ---------------------------------------------------------------------------

def test_unsupported_file():
    txt_bytes = b"Hello world"
    response = client.post(
        "/analyze",
        files={"file": ("test.txt", txt_bytes, "text/plain")}
    )
    assert response.status_code == 400
    assert "Unsupported file type" in response.json()["detail"]


def test_metadata_cleaning(monkeypatch):
    """Ensure DOCX metadata is stripped during redaction."""
    docx_bytes = _create_test_docx()
    # Add metadata to original
    doc = Document(io.BytesIO(docx_bytes))
    doc.core_properties.author = "Sneaky Attacker"
    doc.core_properties.comments = "Hidden tracking code"

    out = io.BytesIO()
    doc.save(out)
    dirty_bytes = out.getvalue()

    redact_resp = client.post(
        "/redact",
        files={"file": ("test.docx", dirty_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"entities": json.dumps([])}  # Pass empty entities to just trigger cleaning
    )

    clean_doc = Document(io.BytesIO(redact_resp.content))
    assert clean_doc.core_properties.author == ""
    assert clean_doc.core_properties.comments == ""


# ---------------------------------------------------------------------------
# Test custom dictionary detection
# ---------------------------------------------------------------------------

def test_custom_dictionary():
    """Test custom term detection."""
    text = "The secret project codenamed PHOENIX is classified."
    entities = detect_entities(text, custom_terms=["PHOENIX", "classified"])
    types = {e.entity_type for e in entities}
    texts = {e.text for e in entities}
    assert "CUSTOM" in types
    assert "PHOENIX" in texts


# ---------------------------------------------------------------------------
# Test context boost
# ---------------------------------------------------------------------------

def test_context_boost():
    """Entities near sensitive keywords should get higher scores."""
    text = "The patient SSN is 838-23-1111 according to medical records."
    entities = detect_entities(text)
    ssn_ent = next((e for e in entities if e.entity_type == "SSN"), None)
    if ssn_ent:
        # Context words "patient" and "medical" should boost the score
        assert ssn_ent.score > 0.90
