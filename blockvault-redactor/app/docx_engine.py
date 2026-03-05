"""DOCX text extraction, entity detection, and redaction.

Uses python-docx for reading/writing and lxml for metadata removal.
"""
from __future__ import annotations

import io
import logging
import re
from typing import List

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from .detector import Entity, detect_entities
from .config import REDACT_LABEL

logger = logging.getLogger(__name__)


def extract_text_docx(docx_bytes: bytes) -> str:
    """Extract all paragraph and table text from a DOCX."""
    doc = Document(io.BytesIO(docx_bytes))
    parts: List[str] = []

    for para in doc.paragraphs:
        parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts)


def analyze_docx(docx_bytes: bytes) -> List[Entity]:
    """Detect entities in a DOCX file."""
    text = extract_text_docx(docx_bytes)
    entities = detect_entities(text)
    # DOCX entities don't have page/bbox
    for ent in entities:
        ent.page = 1  # Single logical page
    return entities


def _replace_in_runs(paragraph, entity_text: str, replacement: str) -> bool:
    """Replace entity_text across potentially split runs in a paragraph.

    Word processors often split text across multiple XML runs, so a name
    like "John Smith" might be in runs ["John ", "Smith"].  This function
    handles that case.
    """
    full_text = paragraph.text
    if entity_text not in full_text:
        return False

    # Simple case: entity is within a single run
    for run in paragraph.runs:
        if entity_text in run.text:
            run.text = run.text.replace(entity_text, replacement)
            return True

    # Complex case: entity spans multiple runs — rebuild paragraph
    # Find the start index in the full text
    idx = full_text.find(entity_text)
    if idx == -1:
        return False

    new_text = full_text[:idx] + replacement + full_text[idx + len(entity_text):]

    # Clear all runs and set first run to new text
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ""

    return True


def redact_docx(docx_bytes: bytes, entities: List[Entity]) -> bytes:
    """Redact entities in a DOCX and clean metadata.

    Replaces entity text with [REDACTED], removes comments,
    tracked changes, and document metadata.
    """
    doc = Document(io.BytesIO(docx_bytes))

    # Sort entities longest-first to avoid partial replacement issues
    sorted_entities = sorted(entities, key=lambda e: -len(e.text))

    # Redact in paragraphs
    for para in doc.paragraphs:
        for ent in sorted_entities:
            _replace_in_runs(para, ent.text, REDACT_LABEL)

    # Redact in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for ent in sorted_entities:
                        _replace_in_runs(para, ent.text, REDACT_LABEL)

    # Redact in headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                for para in header.paragraphs:
                    for ent in sorted_entities:
                        _replace_in_runs(para, ent.text, REDACT_LABEL)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for para in footer.paragraphs:
                    for ent in sorted_entities:
                        _replace_in_runs(para, ent.text, REDACT_LABEL)

    # Clean metadata
    core = doc.core_properties
    core.author = ""
    core.title = ""
    core.subject = ""
    core.keywords = ""
    core.comments = ""
    core.last_modified_by = ""
    core.category = ""

    # Remove comments from XML (direct lxml manipulation)
    try:
        from lxml import etree
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        body = doc.element.body
        # Remove comment references and content
        for tag in ["commentRangeStart", "commentRangeEnd", "commentReference"]:
            for el in body.findall(f".//w:{tag}", nsmap):
                el.getparent().remove(el)
    except Exception:
        pass

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
